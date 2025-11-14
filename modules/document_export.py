"""
Document Export Module for B3PersonalAssistant.

Provides export functionality for academic documents:
- Export to Word (.docx) with formatting
- Export to LaTeX with academic templates
- Export to Markdown
- Export to HTML
"""

import logging
from pathlib import Path
from typing import List, Dict, Optional, Any
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """Represents a document section."""
    title: str
    level: int  # 1=chapter, 2=section, 3=subsection
    content: str
    subsections: List['DocumentSection'] = None


class DocumentExporter:
    """
    Exports documents to various formats for academic writing.

    Supports:
    - Word (DOCX) with academic formatting
    - LaTeX with templates
    - Markdown
    - HTML
    """

    def __init__(self):
        """Initialize document exporter."""
        self.logger = logging.getLogger("document_exporter")

    def export_to_word(
        self,
        content: str,
        output_path: Path,
        title: Optional[str] = None,
        author: Optional[str] = None,
        style: str = "academic"
    ) -> bool:
        """
        Export content to Word document.

        Args:
            content: Content to export (Markdown format)
            output_path: Output file path
            title: Document title
            author: Author name
            style: Document style (academic, simple, report)

        Returns:
            True if successful
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Set margins for academic style
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Add title if provided
            if title:
                title_para = doc.add_heading(title, 0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add author if provided
            if author:
                author_para = doc.add_paragraph(author)
                author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                author_para.runs[0].font.size = Pt(12)

            # Add spacing
            if title or author:
                doc.add_paragraph()

            # Parse and add content
            self._add_markdown_to_docx(doc, content)

            # Save document
            doc.save(output_path)

            self.logger.info(f"Exported to Word: {output_path}")
            return True

        except ImportError:
            self.logger.error("python-docx not installed. Run: pip install python-docx")
            return False
        except Exception as e:
            self.logger.error(f"Error exporting to Word: {e}", exc_info=True)
            return False

    def _add_markdown_to_docx(self, doc, content: str):
        """Parse Markdown and add to Word document."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        lines = content.split('\n')
        in_code_block = False
        code_lines = []

        for line in lines:
            # Code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    code_text = '\n'.join(code_lines)
                    code_para = doc.add_paragraph(code_text)
                    code_para.style = 'List Paragraph'
                    run = code_para.runs[0]
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    code_lines = []
                in_code_block = not in_code_block
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            # Headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)

            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                para = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                para = doc.add_paragraph(line[3:], style='List Number')

            # Bold and italic
            elif line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)

            # Empty line
            else:
                doc.add_paragraph()

    def _add_formatted_text(self, para, text: str):
        """Add text with bold/italic formatting."""
        import re

        # Split by bold and italic markers
        # Simple implementation - can be enhanced
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)

    def export_to_latex(
        self,
        content: str,
        output_path: Path,
        template: str = "article",
        title: Optional[str] = None,
        author: Optional[str] = None,
        bibliography: Optional[str] = None
    ) -> bool:
        """
        Export content to LaTeX document.

        Args:
            content: Content to export (Markdown or plain text)
            output_path: Output file path (.tex)
            template: LaTeX document class (article, report, book)
            title: Document title
            author: Author name
            bibliography: BibTeX bibliography content

        Returns:
            True if successful
        """
        try:
            latex_content = self._generate_latex_document(
                content,
                template,
                title,
                author,
                bibliography
            )

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(latex_content)

            # Also write bibliography if provided
            if bibliography:
                bib_path = output_path.with_suffix('.bib')
                with open(bib_path, 'w', encoding='utf-8') as f:
                    f.write(bibliography)

                self.logger.info(f"Exported bibliography to: {bib_path}")

            self.logger.info(f"Exported to LaTeX: {output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error exporting to LaTeX: {e}", exc_info=True)
            return False

    def _generate_latex_document(
        self,
        content: str,
        template: str,
        title: Optional[str],
        author: Optional[str],
        bibliography: Optional[str]
    ) -> str:
        """Generate complete LaTeX document."""
        import re

        latex = f"\\documentclass[12pt]{{{template}}}\n\n"

        # Packages
        latex += "\\usepackage[utf8]{inputenc}\n"
        latex += "\\usepackage[T1]{fontenc}\n"
        latex += "\\usepackage{amsmath}\n"
        latex += "\\usepackage{graphicx}\n"
        latex += "\\usepackage{hyperref}\n"

        if bibliography:
            latex += "\\usepackage{natbib}\n"

        latex += "\n"

        # Title and author
        if title:
            latex += f"\\title{{{title}}}\n"

        if author:
            latex += f"\\author{{{author}}}\n"

        latex += "\\date{\\today}\n\n"

        # Begin document
        latex += "\\begin{document}\n\n"

        if title:
            latex += "\\maketitle\n\n"

        # Add content
        latex += self._markdown_to_latex(content)

        # Bibliography
        if bibliography:
            latex += "\n\\bibliographystyle{plain}\n"
            latex += "\\bibliography{references}\n"

        # End document
        latex += "\n\\end{document}\n"

        return latex

    def _markdown_to_latex(self, content: str) -> str:
        """Convert Markdown to LaTeX."""
        import re

        latex = ""
        lines = content.split('\n')

        in_code_block = False

        for line in lines:
            # Code blocks
            if line.startswith('```'):
                if in_code_block:
                    latex += "\\end{verbatim}\n"
                else:
                    latex += "\\begin{verbatim}\n"
                in_code_block = not in_code_block
                continue

            if in_code_block:
                latex += line + "\n"
                continue

            # Headings
            if line.startswith('# '):
                latex += f"\\section{{{line[2:]}}}\n"
            elif line.startswith('## '):
                latex += f"\\subsection{{{line[3:]}}}\n"
            elif line.startswith('### '):
                latex += f"\\subsubsection{{{line[4:]}}}\n"

            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                if not latex.endswith('\\begin{itemize}\n'):
                    latex += "\\begin{itemize}\n"
                latex += f"  \\item {line[2:]}\n"
            elif re.match(r'^\d+\. ', line):
                if not latex.endswith('\\begin{enumerate}\n'):
                    latex += "\\begin{enumerate}\n"
                latex += f"  \\item {line[3:]}\n"

            # Regular text
            elif line.strip():
                # Close lists if open
                if latex.endswith('\\begin{itemize}\n'):
                    latex = latex[:-17]  # Remove \begin{itemize}
                elif latex.endswith('\\begin{enumerate}\n'):
                    latex = latex[:-19]

                # Convert bold and italic
                line = re.sub(r'\*\*(.*?)\*\*', r'\\textbf{\1}', line)
                line = re.sub(r'\*(.*?)\*', r'\\textit{\1}', line)

                latex += line + "\n"

            # Empty line
            else:
                latex += "\n"

        return latex

    def export_to_markdown(
        self,
        sections: List[DocumentSection],
        output_path: Path,
        title: Optional[str] = None
    ) -> bool:
        """
        Export structured sections to Markdown.

        Args:
            sections: List of DocumentSection objects
            output_path: Output file path
            title: Document title

        Returns:
            True if successful
        """
        try:
            markdown = ""

            if title:
                markdown += f"# {title}\n\n"

            for section in sections:
                markdown += self._section_to_markdown(section)

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown)

            self.logger.info(f"Exported to Markdown: {output_path}")
            return True

        except Exception as e:
            self.logger.error(f"Error exporting to Markdown: {e}", exc_info=True)
            return False

    def _section_to_markdown(self, section: DocumentSection, base_level: int = 0) -> str:
        """Convert section to Markdown."""
        level = section.level + base_level
        heading = "#" * level + " " + section.title + "\n\n"

        content = section.content + "\n\n"

        subsections = ""
        if section.subsections:
            for subsection in section.subsections:
                subsections += self._section_to_markdown(subsection, base_level)

        return heading + content + subsections


def create_document_exporter() -> DocumentExporter:
    """Convenience function to create document exporter."""
    return DocumentExporter()


# Import for formatted text parsing
import re
